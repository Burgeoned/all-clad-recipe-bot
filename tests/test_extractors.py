import io

import pytest
from docx import Document

from recipe_pipeline.extractors import ExtractionError, extract_text
from recipe_pipeline.models import FileType, SourceFile


def _source(name: str, file_type: FileType, content: bytes) -> SourceFile:
    return SourceFile(file_id="id", name=name, file_type=file_type, content=content)


def test_extract_txt_utf8():
    text = extract_text(_source("r.txt", FileType.TXT, "héllo wörld".encode("utf-8")))
    assert "héllo wörld" in text


def test_extract_txt_falls_back_to_cp1252():
    # A byte sequence that is invalid UTF-8 but valid cp1252.
    text = extract_text(_source("r.txt", FileType.TXT, "café".encode("cp1252")))
    assert "café" in text


def test_extract_docx_reads_paragraphs_and_tables():
    doc = Document()
    doc.add_paragraph("Chicken Soup")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "2 cups"
    table.rows[0].cells[1].text = "broth"
    buffer = io.BytesIO()
    doc.save(buffer)

    text = extract_text(_source("r.docx", FileType.DOCX, buffer.getvalue()))
    assert "Chicken Soup" in text
    assert "2 cups" in text
    assert "broth" in text


def test_extract_pdf_unreadable_raises():
    with pytest.raises(ExtractionError):
        extract_text(_source("r.pdf", FileType.PDF, b"not a real pdf"))
