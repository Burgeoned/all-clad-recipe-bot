"""Turn a downloaded recipe file into plain text for the parser.

Extraction is deliberately dumb: pull *all* the text out, in reading order, and let Claude
make sense of structure downstream. We don't try to interpret layout here.
"""

from __future__ import annotations

import io

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from .models import FileType, SourceFile


class ExtractionError(Exception):
    """Raised when a source file's text cannot be extracted."""


def extract_text(source: SourceFile) -> str:
    """Dispatch on file type and return the file's text content.

    Raises ExtractionError on unreadable/corrupt files or an unsupported type.
    """
    if source.file_type is FileType.TXT:
        return _extract_txt(source.content)
    if source.file_type is FileType.DOCX:
        return _extract_docx(source.content)
    if source.file_type is FileType.PDF:
        return _extract_pdf(source.content)
    raise ExtractionError(f"Unsupported file type: {source.file_type}")


def _extract_txt(content: bytes) -> str:
    # Recipes are usually UTF-8; fall back to cp1252 (common from Windows/Word exports)
    # rather than crashing on a stray byte.
    for encoding in ("utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode text file with utf-8, cp1252, or latin-1")


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except (ValueError, KeyError, OSError) as exc:
        # python-docx raises these on a non-docx / corrupt package.
        raise ExtractionError(f"Could not open .docx file: {exc}") from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    # Recipes are frequently laid out in tables; capture cell text too.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except (PdfReadError, ValueError, OSError) as exc:
        raise ExtractionError(f"Could not read PDF file: {exc}") from exc

    text = "\n".join(pages).strip()
    if not text:
        # Almost always a scanned/image-only PDF — OCR is out of scope for v1.
        raise ExtractionError(
            "No extractable text (likely a scanned/image-only PDF; OCR is not supported)"
        )
    return text
