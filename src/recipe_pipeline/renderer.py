"""Fill Recipe_Template.docx from a Recipe, and optionally export it to PDF.

The template is filled in place by **structural anchors** — heading text and table headers —
never by injecting tokens, so the Drive template stays pristine. See DESIGN.md §4.2.

Filling strategy:
  * scalar slots (title, description, metadata lines) → replace the placeholder paragraph's
    text, preserving its run formatting;
  * list sections (substitutions, equipment, steps, tips) → reuse the placeholder list
    paragraphs, cloning the last one for extra items and deleting unused ones, so the list
    style / numbering is preserved;
  * the three data tables (at-a-glance, ingredients, nutrition) → write into the data rows,
    cloning/removing rows in the ingredients table to match the ingredient count.
"""

from __future__ import annotations

import copy
import io
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .config import Settings
from .models import Recipe, RenderedDocs


class RenderError(Exception):
    """Raised when the template cannot be filled or exported."""


EM_DASH = "—"


def render(recipe: Recipe, template_bytes: bytes, settings: Settings) -> RenderedDocs:
    """Fill the template with `recipe`; export to PDF when settings.export_pdf is on."""
    try:
        document = Document(io.BytesIO(template_bytes))
    except (ValueError, KeyError, OSError) as exc:
        raise RenderError(f"Could not open template docx: {exc}") from exc

    _fill_header(document, recipe)
    _fill_metadata(document, recipe)
    _fill_list_sections(document, recipe)
    _fill_tables(document, recipe)

    buffer = io.BytesIO()
    document.save(buffer)
    docx_bytes = buffer.getvalue()

    pdf_bytes = _export_pdf(docx_bytes) if settings.export_pdf else None
    return RenderedDocs(docx=docx_bytes, pdf=pdf_bytes)


# --------------------------------------------------------------------------------------
# Scalar slots
# --------------------------------------------------------------------------------------


def _fill_header(document: DocxDocument, recipe: Recipe) -> None:
    name_para = _find_paragraph(document, lambda t: t == "[ Recipe Name ]")
    if name_para is not None:
        _set_text(name_para, recipe.title)

    desc_para = _find_paragraph(document, lambda t: t.startswith("A short, mouth-watering"))
    if desc_para is not None:
        if recipe.description:
            _set_text(desc_para, recipe.description)
        else:
            _delete_paragraph(desc_para)


def _fill_metadata(document: DocxDocument, recipe: Recipe) -> None:
    # Each metadata line is "Label: placeholder". Rewrite as "Label: value", or drop the
    # line when we have no value.
    fields: list[tuple[str, str | None]] = [
        ("Cuisine", recipe.cuisine),
        ("Course", recipe.course),
        ("Difficulty", recipe.difficulty),
        ("Dietary Tags", ", ".join(recipe.dietary_tags) if recipe.dietary_tags else None),
        ("Source", recipe.source),
    ]
    for label, value in fields:
        para = _find_paragraph(document, lambda t, lbl=label: t.startswith(f"{lbl}:"))
        if para is None:
            continue
        if value:
            _set_text(para, f"{label}: {value}")
        else:
            _delete_paragraph(para)


# --------------------------------------------------------------------------------------
# List sections
# --------------------------------------------------------------------------------------


def _fill_list_sections(document: DocxDocument, recipe: Recipe) -> None:
    _fill_list_between(document, "Substitutions & Notes", "EQUIPMENT", _substitutions_and_notes(recipe))
    _fill_list_between(document, "EQUIPMENT", "INSTRUCTIONS", recipe.equipment)
    _fill_list_between(document, "INSTRUCTIONS", "TIPS & VARIATIONS", recipe.steps)
    _fill_list_between(document, "TIPS & VARIATIONS", "NUTRITION", recipe.tips)


def _substitutions_and_notes(recipe: Recipe) -> list[str]:
    """The 'Substitutions & Notes' block also carries the estimation marker and any warnings,
    so they're visible in the doc rather than buried in a log."""
    notes = list(recipe.substitutions)
    if recipe.nutrition_estimated:
        notes.append("Nutrition is estimated from the ingredients (approximate).")
    notes.extend(f"⚠ {warning}" for warning in recipe.warnings)
    return notes


def _fill_list_between(
    document: DocxDocument, start_heading: str, end_heading: str, items: list[str]
) -> None:
    """Replace the placeholder list paragraphs between two headings with `items`."""
    paragraphs = document.paragraphs
    start = _index_of(paragraphs, lambda t: t.upper() == start_heading.upper())
    end = _index_of(paragraphs, lambda t: t.upper().startswith(end_heading.upper()))
    if start is None or end is None or end <= start:
        return

    placeholders = [p for p in paragraphs[start + 1 : end] if p.text.strip()]
    if not placeholders:
        return

    if not items:
        # Nothing to show: keep one blank so the section doesn't collapse oddly, drop rest.
        _set_text(placeholders[0], EM_DASH)
        for extra in placeholders[1:]:
            _delete_paragraph(extra)
        return

    template_para = placeholders[-1]
    for i, item in enumerate(items):
        if i < len(placeholders):
            _set_text(placeholders[i], item)
        else:
            new_para = _clone_paragraph_after(template_para)
            _set_text(new_para, item)
            template_para = new_para

    # Remove any leftover placeholders we didn't reuse.
    for extra in placeholders[len(items):]:
        _delete_paragraph(extra)


# --------------------------------------------------------------------------------------
# Tables
# --------------------------------------------------------------------------------------


def _fill_tables(document: DocxDocument, recipe: Recipe) -> None:
    glance = _find_table(document, ["PREP TIME", "COOK TIME", "TOTAL TIME", "SERVINGS"])
    if glance is not None:
        _set_row(glance, 1, [
            _minutes(recipe.prep_time_min),
            _minutes(recipe.cook_time_min),
            _minutes(recipe.total_time_min),
            recipe.servings or EM_DASH,
        ])

    ingredients = _find_table(document, ["QUANTITY", "UNIT", "INGREDIENT", "NOTES"])
    if ingredients is not None:
        _fill_ingredients(ingredients, recipe)

    nutrition = _find_table(
        document, ["CALORIES", "PROTEIN", "CARBS", "FAT", "FIBER", "SODIUM"]
    )
    if nutrition is not None and recipe.nutrition is not None:
        n = recipe.nutrition
        _set_row(nutrition, 1, [
            n.calories or EM_DASH, n.protein or EM_DASH, n.carbs or EM_DASH,
            n.fat or EM_DASH, n.fiber or EM_DASH, n.sodium or EM_DASH,
        ])


def _fill_ingredients(table: Table, recipe: Recipe) -> None:
    rows: list[tuple[str, str, str, str]] = []
    multi = len(recipe.ingredient_groups) > 1
    for group in recipe.ingredient_groups:
        if group.heading and multi:
            # Sub-component label as its own row, name in the INGREDIENT column.
            rows.append(("", "", group.heading, ""))
        for ing in group.ingredients:
            rows.append((ing.quantity or "", ing.unit or "", ing.item, ing.note or ""))

    if not rows:
        return

    template_tr = table.rows[1]._tr  # first data row carries the cell styling
    for i, cells in enumerate(rows):
        row_index = 1 + i
        if row_index < len(table.rows):
            row = table.rows[row_index]
        else:
            new_tr = copy.deepcopy(template_tr)
            table._tbl.append(new_tr)
            row = table.rows[-1]
        for c, value in enumerate(cells):
            _set_cell(row.cells[c], value)

    # Drop leftover blank template rows.
    while len(table.rows) > 1 + len(rows):
        last = table.rows[-1]
        last._tr.getparent().remove(last._tr)


def _set_row(table: Table, row_index: int, values: list[str]) -> None:
    row = table.rows[row_index]
    for cell, value in zip(row.cells, values):
        _set_cell(cell, value)


# --------------------------------------------------------------------------------------
# PDF export
# --------------------------------------------------------------------------------------


def _export_pdf(docx_bytes: bytes) -> bytes:
    """Convert docx → pdf via headless LibreOffice. Raises RenderError if unavailable."""
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if soffice is None:
        raise RenderError(
            "export_pdf is on but LibreOffice (soffice) was not found on PATH. "
            "Install LibreOffice or set EXPORT_PDF=0."
        )

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        src = tmp_path / "recipe.docx"
        src.write_bytes(docx_bytes)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path),
                 str(src)],
                check=True, capture_output=True, timeout=120,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise RenderError(f"LibreOffice PDF conversion failed: {exc}") from exc

        pdf = tmp_path / "recipe.pdf"
        if not pdf.is_file():
            raise RenderError("LibreOffice ran but produced no PDF.")
        return pdf.read_bytes()


# --------------------------------------------------------------------------------------
# docx helpers
# --------------------------------------------------------------------------------------


def _minutes(value: int | None) -> str:
    return f"{value} min" if value is not None else EM_DASH


def _set_text(paragraph: Paragraph, text: str) -> None:
    """Set a paragraph's text while keeping the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _set_cell(cell: _Cell, text: str) -> None:
    _set_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def _delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._p.getparent().remove(paragraph._p)


def _clone_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _find_paragraph(document: DocxDocument, predicate) -> Paragraph | None:
    for para in document.paragraphs:
        if predicate(para.text.strip()):
            return para
    return None


def _index_of(paragraphs, predicate) -> int | None:
    for i, para in enumerate(paragraphs):
        if predicate(para.text.strip()):
            return i
    return None


def _find_table(document: DocxDocument, header: list[str]) -> Table | None:
    want = [h.upper() for h in header]
    for table in document.tables:
        if not table.rows:
            continue
        first = [c.text.strip().upper() for c in table.rows[0].cells]
        if first[: len(want)] == want:
            return table
    return None
